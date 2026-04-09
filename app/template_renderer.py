from __future__ import annotations

import os
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document


@dataclass
class RenderPaths:
    docx_path: str
    pdf_path: str


REPLACEMENTS = {
    "North2026_NOR1-0117": "{DOC_NO}",
    "01 เมษายน 2569": "{DOC_DATE_THAI}",
    "เจ้าของอาคาร / เจ้าของพื้นที่": "{RECIPIENT}",
    "CMIA218": "{SITE_CODE}",
    "ตรวจสอบแก้ไขระบบไฟฟ้า": "{WORK_DETAIL}",
    "01 – 03 เมษายน พ.ศ. 2569": "{WORK_DATE}",
    "08:00 – 17:30 น.": "{WORK_TIME}",
    "นายศรัณย์ พรหมอนันต์": "{STAFF_1_NAME}",
    "095-778-6275": "{STAFF_1_PHONE}",
    "นายกฤษณะ ปอเย็น": "{STAFF_2_NAME}",
    "นายจตุพล จันทร์ทอง": "{STAFF_3_NAME}",
    "081-595-2897": "{STAFF_3_PHONE}",
    "ผู้ดำเนินงาน BBTEC": "{STAFF_3_ROLE}",
    "นายพิพัฒน์พงษ์ ศึกษา": "{STAFF_4_NAME}",
    "086-522-0545": "{STAFF_4_PHONE}",
    "ผู้ควบคุมงาน TRUE": "{STAFF_4_ROLE}",
}


def create_placeholder_template(src_docx: str, dest_docx: str) -> str:
    doc = Document(src_docx)
    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        if not full_text:
            continue
        updated = full_text
        for old, new in REPLACEMENTS.items():
            updated = updated.replace(old, new)
        if updated != full_text:
            if paragraph.runs:
                paragraph.runs[0].text = updated
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.text = updated
    doc.save(dest_docx)
    return dest_docx


def _fill_staff(staff: list[dict[str, str]]) -> dict[str, str]:
    result: dict[str, str] = {}
    for i in range(4):
        row = staff[i] if i < len(staff) else {}
        idx = i + 1
        result[f"STAFF_{idx}_NAME"] = row.get("name", "")
        result[f"STAFF_{idx}_PHONE"] = row.get("phone", "")
        result[f"STAFF_{idx}_ROLE"] = row.get("role", "")
    return result


def fill_template(template_docx: str, output_docx: str, data: dict[str, str], staff: list[dict[str, str]]) -> str:
    doc = Document(template_docx)
    flat_data = dict(data)
    flat_data.update(_fill_staff(staff))
    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        if not full_text:
            continue
        updated = full_text
        for key, value in flat_data.items():
            updated = updated.replace(f"{{{key}}}", str(value or ""))
        if updated != full_text:
            if paragraph.runs:
                paragraph.runs[0].text = updated
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.text = updated
    doc.save(output_docx)
    return output_docx


def convert_docx_to_pdf(docx_path: str, output_dir: str) -> str:
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not libreoffice:
        raise RuntimeError("LibreOffice/soffice not found in PATH")
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    cmd = [
        libreoffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        output_dir,
        docx_path,
    ]
    proc = subprocess.run(cmd, capture_output=True, text=True)
    if proc.returncode != 0:
        raise RuntimeError(f"PDF conversion failed: {proc.stderr or proc.stdout}")
    pdf_path = str(Path(output_dir) / (Path(docx_path).stem + ".pdf"))
    if not os.path.exists(pdf_path):
        raise RuntimeError("PDF file not generated")
    return pdf_path
